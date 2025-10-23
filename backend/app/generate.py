"""
Document generation utilities for filling DOCX templates and HTML previews
"""
import re
from docx import Document
from typing import Dict, Set

# Regex pattern to match [[TOKEN]] placeholders (double bracket)
TOKEN = re.compile(r"\[\[([A-Z0-9_]+)\]\]")

# Also support single-bracket legacy format [TOKEN]
LEGACY_TOKEN = re.compile(r"\[([A-Z0-9_ ]+)\]")

def _replace_runs(paragraph, replacements: Dict[str, str]) -> Set[str]:
    """
    Replace tokens in a paragraph while preserving formatting.
    
    Word documents often split text across multiple "runs" (text fragments with
    different formatting). We merge all runs, replace tokens, then put the result
    back in the first run to maintain formatting.
    
    Supports both [[TOKEN]] and legacy [TOKEN] formats.
    
    Args:
        paragraph: python-docx Paragraph object
        replacements: Dict mapping token names (uppercase) to replacement values
    
    Returns:
        Set of tokens that were found but had no replacement value
    """
    # Merge all run text to handle tokens split across runs
    text = "".join(r.text for r in paragraph.runs)
    missing_tokens = set()
    
    # Track which tokens were found but not replaced
    def replace_token(match):
        token = match.group(1).strip()
        if token in replacements:
            return str(replacements[token])
        else:
            missing_tokens.add(token)
            return match.group(0)  # Leave unchanged if no replacement
    
    # Replace [[TOKEN]] patterns
    new = TOKEN.sub(replace_token, text)
    
    # Also handle legacy [TOKEN] format (single brackets)
    new = LEGACY_TOKEN.sub(replace_token, new)
    
    # Remove all runs except the first one
    for _ in range(len(paragraph.runs) - 1):
        paragraph.runs[0]._element.getparent().remove(paragraph.runs[1]._element)
    
    # Put the replaced text in the first run (preserves its formatting)
    if paragraph.runs:
        paragraph.runs[0].text = new
    
    return missing_tokens

def fill_docx(input_path: str, output_path: str, replacements: Dict[str, str]) -> Set[str]:
    """
    Fill a DOCX template with replacement values.
    
    Searches through all paragraphs and table cells for [[TOKEN]] patterns
    and replaces them with values from the replacements dict.
    
    Args:
        input_path: Path to template DOCX file
        output_path: Path to save filled DOCX file
        replacements: Dict mapping token names (uppercase) to replacement values
    
    Returns:
        Set of tokens that were found but not replaced (missing values)
    """
    doc = Document(input_path)
    missing_tokens = set()
    
    # Replace tokens in paragraphs
    for p in doc.paragraphs:
        text = "".join(r.text for r in p.runs)
        if TOKEN.search(text) or LEGACY_TOKEN.search(text):
            missing = _replace_runs(p, replacements)
            missing_tokens.update(missing)
    
    # Replace tokens in table cells
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    text = "".join(r.text for r in p.runs)
                    if TOKEN.search(text) or LEGACY_TOKEN.search(text):
                        missing = _replace_runs(p, replacements)
                        missing_tokens.update(missing)
    
    doc.save(output_path)
    return missing_tokens

def render_html(html_template: str, replacements: Dict[str, str]) -> str:
    """
    Fill HTML template and wrap replaced values in <mark> tags for review.
    
    This creates a highlighted preview where lawyers can easily see what
    values were filled in by the client.
    
    Args:
        html_template: HTML string with [[TOKEN]] placeholders
        replacements: Dict mapping token names (uppercase) to replacement values
    
    Returns:
        HTML string with tokens replaced and wrapped in <mark data-key="TOKEN">
    """
    return TOKEN.sub(
        lambda m: f'<mark data-key="{m.group(1)}">{replacements.get(m.group(1), m.group(0))}</mark>',
        html_template
    )

